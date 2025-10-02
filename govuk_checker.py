# govuk_checker.py
# A Python script to check HTM, HTML, and DOCX files for GOV.UK style guide compliance and specified content criteria.
# Usage:
#   python govuk_checker.py <file_path>
# Supported file types: .htm, .html, .docx

import sys
import os
import re
from bs4 import BeautifulSoup
from docx import Document

def check_html(file_path):
    findings = []
    with open(file_path, "r", encoding="utf-8") as f:
        soup = BeautifulSoup(f, "lxml")
    text = soup.get_text(separator="\n")

    bullet_points = [li.get_text() for li in soup.find_all("li")]
    for bp in bullet_points:
        if not bp.strip().endswith("."):
            findings.append(f"BULLET: Does not end with a full stop: '{bp.strip()}'")
        if bp and not bp.strip()[0].isupper():
            findings.append(f"BULLET: Does not begin with a capital letter: '{bp.strip()}'")

    for ul in soup.find_all("ul"):
        prev = ul.find_previous(string=True)
        if not prev or not prev.strip().endswith(":"):
            findings.append("BULLET: List may be missing a lead-in sentence.")

    if re.search(r"\bplease\b", text, re.IGNORECASE):
        findings.append("LANGUAGE: Found instance of the word 'please'.")

    neg_contr = re.findall(r"\b(?:don’t|doesn’t|didn’t|can’t|won’t|wouldn’t|shouldn’t|couldn’t|isn’t|aren’t|wasn’t|weren’t|haven’t|hasn’t|hadn’t|mustn’t|mightn’t|needn’t)\b", text, re.IGNORECASE)
    for nc in neg_contr:
        findings.append(f"LANGUAGE: Found negative contraction: '{nc}'")

    for word in ['above', 'below']:
        if re.search(rf"\b{word}\b", text, re.IGNORECASE):
            findings.append(f"LANGUAGE: Found use of the word '{word}'.")

    sentences = re.split(r'(?<=[.!?])\s+', text)
    for sentence in sentences:
        if len(sentence.split()) > 26:
            findings.append(f"LANGUAGE: Long sentence (>26 words): '{sentence.strip()}'")

    for a in soup.find_all("a"):
        link_text = a.get_text().strip()
        if len(link_text.split()) == 1:
            findings.append(f"LINK: Link text is only one word: '{link_text}'")
        if not re.search(r"\b(work|read|learn|report|check|view|explore|download|submit|apply|contact|visit)\b", link_text, re.IGNORECASE):
            findings.append(f"LINK: Link text may not be descriptive or active: '{link_text}'")

    images = soup.find_all("img")
    for img in images:
        alt_text = img.get("alt", "")
        if not alt_text:
            findings.append("IMAGE: Image found without alt text.")
        if alt_text and alt_text not in text:
            findings.append(f"IMAGE: Alt text not described in body: '{alt_text}'")

    for td in soup.find_all("td"):
        if not td.get_text(strip=True):
            findings.append("TABLE: Empty table cell found. Should be marked as 'no data' or 'not applicable'.")

    for tag in soup.find_all(["b", "strong", "i", "em"]):
        findings.append(f"FORMAT: Use of bold/italic text: '{tag.get_text(strip=True)}'")

    acronyms = re.findall(r"\b[A-Z]{2,}\b", text)
    for acronym in set(acronyms):
        if not re.search(rf"\b{acronym} \([^)]+\)", text):
            findings.append(f"ACRONYM: '{acronym}' may not be spelled out on first use.")

    if not re.search(r"\b\d{1,2} (January|February|March|April|May|June|July|August|September|October|November|December) \d{4}\b", text):
        findings.append("STYLE: No date found in GOV.UK format (e.g., '21 September 2025').")
    if re.search(r"\s-\s", text):
        findings.append("STYLE: Hyphen used where em dash may be appropriate.")
    if "/" in text:
        findings.append("STYLE: Slash '/' found in text.")
    for table in soup.find_all("table"):
        if table.find("a"):
            findings.append("STYLE: Link found inside a table.")
    titles = soup.find_all(["h1", "h2", "h3"])
    for title in titles:
        if not re.search(r"\b(work|read|learn|report|check|view|explore|download|submit|apply|contact|visit)\b", title.get_text(), re.IGNORECASE):
            findings.append(f"STYLE: Title may not use active language: '{title.get_text(strip=True)}'")

    return findings

def check_docx(file_path):
    findings = []
    doc = Document(file_path)
    text = "\n".join([p.text for p in doc.paragraphs])

    for para in doc.paragraphs:
        if para.style.name.startswith("List"):
            bp = para.text.strip()
            if not bp.endswith("."):
                findings.append(f"BULLET: Does not end with a full stop: '{bp}'")
            if bp and not bp[0].isupper():
                findings.append(f"BULLET: Does not begin with a capital letter: '{bp}'")

    if re.search(r"\bplease\b", text, re.IGNORECASE):
        findings.append("LANGUAGE: Found instance of the word 'please'.")

    neg_contr = re.findall(r"\b(?:don’t|doesn’t|didn’t|can’t|won’t|wouldn’t|shouldn’t|couldn’t|isn’t|aren’t|wasn’t|weren’t|haven’t|hasn’t|hadn’t|mustn’t|mightn’t|needn’t)\b", text, re.IGNORECASE)
    for nc in neg_contr:
        findings.append(f"LANGUAGE: Found negative contraction: '{nc}'")

    for word in ['above', 'below']:
        if re.search(rf"\b{word}\b", text, re.IGNORECASE):
            findings.append(f"LANGUAGE: Found use of the word '{word}'.")

    sentences = re.split(r'(?<=[.!?])\s+', text)
    for sentence in sentences:
        if len(sentence.split()) > 26:
            findings.append(f"LANGUAGE: Long sentence (>26 words): '{sentence.strip()}'")

    for para in doc.paragraphs:
        for run in para.runs:
            if run.bold or run.italic:
                findings.append(f"FORMAT: Use of bold/italic text: '{run.text.strip()}'")

    acronyms = re.findall(r"\b[A-Z]{2,}\b", text)
    for acronym in set(acronyms):
        if not re.search(rf"\b{acronym} \([^)]+\)", text):
            findings.append(f"ACRONYM: '{acronym}' may not be spelled out on first use.")

    if not re.search(r"\b\d{1,2} (January|February|March|April|May|June|July|August|September|October|November|December) \d{4}\b", text):
        findings.append("STYLE: No date found in GOV.UK format (e.g., '21 September 2025').")
    if re.search(r"\s-\s", text):
        findings.append("STYLE: Hyphen used where em dash may be appropriate.")
    if "/" in text:
        findings.append("STYLE: Slash '/' found in text.")

    return findings

def main():
    if len(sys.argv) != 2:
        print("Usage: python govuk_checker.py <file>")
        return
    file_path = sys.argv[1]
    if not os.path.isfile(file_path):
        print("File not found.")
        return
    ext = os.path.splitext(file_path)[1].lower()
    if ext in [".htm", ".html"]:
        findings = check_html(file_path)
    elif ext == ".docx":
        findings = check_docx(file_path)
    else:
        print("Unsupported file type. Use .htm, .html, or .docx")
        return
    print("\n--- GOV.UK Style & Content Findings ---\n")
    for f in findings:
        print(f)
    print(f"\nTotal findings: {len(findings)}")

if __name__ == "__main__":
    main()
